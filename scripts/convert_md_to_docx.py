import os
import re
from docx import Document

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        # Headings
        if line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        # Tables
        elif '|' in line and i + 1 < len(lines) and set(lines[i+1].strip()) <= set('-| '):
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].rstrip('\n'))
                i += 1
            headers = [h.strip() for h in table_lines[0].split('|')[1:-1]]
            rows = []
            for row in table_lines[2:]:
                rows.append([c.strip() for c in row.split('|')[1:-1]])
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            # Add header cells
            for idx, head in enumerate(headers):
                table.rows[0].cells[idx].text = head
            # Add data rows
            for r, data in enumerate(rows, start=1):
                for c, cell in enumerate(data):
                    table.rows[r].cells[c].text = cell
            continue
        # Paragraphs with bold emphasis
        else:
            para = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                m = re.match(r'\*\*(.*?)\*\*', part)
                if m:
                    run = para.add_run(m.group(1))
                    run.bold = True
                else:
                    para.add_run(part)
        i += 1
    # Ensure output directory exists
    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)

def main():
    src_dir = 'report'
    for root, _, files in os.walk(src_dir):
        for fname in files:
            if fname.endswith('.md'):
                md_path = os.path.join(root, fname)
                rel = os.path.relpath(root, src_dir)
                out_dir = os.path.join(src_dir, rel)
                base = os.path.splitext(fname)[0]
                docx_path = os.path.join(out_dir, f'{base}.docx')
                convert_md_to_docx(md_path, docx_path)

if __name__ == '__main__':
    main()
